"""Build the scientific report figures and a polished DOCX.

Uses the bundled document runtime (python-docx + Pillow). The report text is
maintained in reports/scientific_report.md and results are read from runs/.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
SOURCE = ROOT / 'reports/scientific_report.md'
FIGURES = ROOT / 'reports/figures'
OUTPUT = ROOT / 'output/docs/FinanceBench_RAG_Scientific_Report.docx'

SKILL_SCRIPTS = Path(
    '/Users/nayeli/.codex/plugins/cache/openai-primary-runtime/documents/'
    '26.723.12215/skills/documents/scripts'
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


# narrative_proposal preset with a named scientific-report override:
# Arial body for cross-platform availability; editorial title in Georgia.
NAVY = '#17324D'
BLUE = '#2E5D7B'
TEAL = '#2C7A7B'
ORANGE = '#D97706'
SLATE = '#64748B'
LIGHT = '#EEF3F6'
PALE = '#F7F9FB'
GREEN = '#3B806B'
RED = '#B64A4A'
BLACK = '#17212B'
WHITE = '#FFFFFF'

FONT_REG = '/System/Library/Fonts/Supplemental/Arial.ttf'
FONT_BOLD = '/System/Library/Fonts/Supplemental/Arial Bold.ttf'
FONT_SERIF = '/System/Library/Fonts/Supplemental/Georgia.ttf'
FONT_SERIF_BOLD = '/System/Library/Fonts/Supplemental/Georgia Bold.ttf'


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip('#')
    return RGBColor.from_string(value)


def pil_font(size: int, bold: bool = False, serif: bool = False):
    path = FONT_SERIF_BOLD if (serif and bold) else (
        FONT_SERIF if serif else (FONT_BOLD if bold else FONT_REG)
    )
    return ImageFont.truetype(path, size)


def center_text(draw, xy, text, font, fill, anchor='mm'):
    draw.text(xy, text, font=font, fill=fill, anchor=anchor)


def save_canvas(image: Image.Image, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, format='PNG', optimize=True)


def figure_pipeline():
    w, h = 1800, 760
    image = Image.new('RGB', (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    title = pil_font(42, bold=True, serif=True)
    subtitle = pil_font(23)
    label = pil_font(24, bold=True)
    detail = pil_font(18)
    center_text(draw, (w / 2, 55), 'Closed-Corpus Financial RAG Pipeline',
                title, NAVY)
    center_text(draw, (w / 2, 100),
                'From source filing to evaluated, cited answer',
                subtitle, SLATE)

    boxes = [
        (110, 185, 390, 315, '368 PDFs', 'FinanceBench corpus', BLUE),
        (455, 185, 735, 315, '54,120 pages', 'Parse + clean', TEAL),
        (800, 185, 1080, 315, '211,671 chunks', 'Page-bounded', ORANGE),
        (1145, 185, 1425, 315, '384-d vectors', 'BGE-small', BLUE),
        (490, 455, 800, 585, 'Exact retrieval', 'FAISS + optional BM25', TEAL),
        (875, 455, 1185, 585, 'Local generation', 'Qwen2.5 + citations', ORANGE),
        (1260, 455, 1570, 585, 'Evaluation', 'Answer + evidence', BLUE),
    ]
    for x1, y1, x2, y2, head, sub, color in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20,
                               fill=PALE, outline=color, width=5)
        center_text(draw, ((x1 + x2) / 2, y1 + 47), head, label, color)
        center_text(draw, ((x1 + x2) / 2, y1 + 91), sub, detail, BLACK)

    def arrow(start, end):
        draw.line((start, end), fill=SLATE, width=6)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) > abs(ey - sy):
            sign = 1 if ex > sx else -1
            pts = [(ex, ey), (ex - sign * 20, ey - 12),
                   (ex - sign * 20, ey + 12)]
        else:
            sign = 1 if ey > sy else -1
            pts = [(ex, ey), (ex - 12, ey - sign * 20),
                   (ex + 12, ey - sign * 20)]
        draw.polygon(pts, fill=SLATE)

    arrow((390, 250), (455, 250))
    arrow((735, 250), (800, 250))
    arrow((1080, 250), (1145, 250))
    arrow((1285, 315), (645, 455))
    arrow((800, 520), (875, 520))
    arrow((1185, 520), (1260, 520))
    center_text(draw, (w / 2, 685),
                'Every chunk retains document, page, text, and configuration provenance',
                pil_font(21), SLATE)
    save_canvas(image, FIGURES / 'figure_1_pipeline.png')


def draw_panel_axes(draw, box, title, y_label):
    x1, y1, x2, y2 = box
    draw.text((x1, y1 - 55), title, font=pil_font(28, bold=True), fill=NAVY)
    plot = (x1 + 80, y1 + 20, x2 - 25, y2 - 70)
    px1, py1, px2, py2 = plot
    draw.line((px1, py1, px1, py2), fill=SLATE, width=3)
    draw.line((px1, py2, px2, py2), fill=SLATE, width=3)
    for i in range(6):
        value = i / 10
        y = py2 - (value / 0.5) * (py2 - py1)
        draw.line((px1, y, px2, y), fill='#D9E1E7', width=2)
        draw.text((px1 - 15, y), f'{value:.1f}', font=pil_font(16),
                  fill=SLATE, anchor='rm')
    draw.text((x1 + 8, (py1 + py2) / 2), y_label, font=pil_font(17),
              fill=SLATE, anchor='mm')
    return plot


def figure_retrieval():
    dense = pd.read_csv(ROOT / 'runs/retrieval_bge-small__s256_o50.csv',
                        index_col=0).loc['overall']
    hybrid = pd.read_csv(
        ROOT / 'runs/retrieval_hybrid-d0p90__bge-small__s256_o50.csv',
        index_col=0).loc['overall']
    ks = [1, 3, 5, 10, 20]
    w, h = 1800, 920
    image = Image.new('RGB', (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    center_text(draw, (w / 2, 55), 'Retrieval Quality by Candidate Depth',
                pil_font(42, bold=True, serif=True), NAVY)
    center_text(draw, (w / 2, 100), 'Full 150-question evaluation',
                pil_font(22), SLATE)
    panels = [
        ((80, 205, 865, 790), 'Gold-page retrieval', 'Page hit'),
        ((935, 205, 1720, 790), 'Exact evidence-text retrieval', 'Evidence hit'),
    ]
    for panel, title, ylabel in panels:
        plot = draw_panel_axes(draw, panel, title, ylabel)
        px1, py1, px2, py2 = plot
        xs = [px1 + i * (px2 - px1) / (len(ks) - 1) for i in range(len(ks))]
        for x, k in zip(xs, ks):
            draw.text((x, py2 + 18), str(k), font=pil_font(18),
                      fill=BLACK, anchor='ma')
        draw.text(((px1 + px2) / 2, py2 + 50), 'Retrieved passages (k)',
                  font=pil_font(18), fill=SLATE, anchor='ma')
        metric = 'page_hit' if 'page' in title.lower() else 'evidence_hit'
        for series, color, name in [
            (dense, BLUE, 'Dense'),
            (hybrid, ORANGE, 'Hybrid 90/10'),
        ]:
            vals = [float(series[f'{metric}@{k}']) for k in ks]
            pts = [(x, py2 - (v / 0.5) * (py2 - py1))
                   for x, v in zip(xs, vals)]
            draw.line(pts, fill=color, width=6)
            for x, y in pts:
                draw.ellipse((x - 8, y - 8, x + 8, y + 8),
                             fill=WHITE, outline=color, width=5)
            draw.text((pts[-1][0] - 6, pts[-1][1] - 22),
                      f'{vals[-1]:.3f}', font=pil_font(17, bold=True),
                      fill=color, anchor='rs')
    legend_y = 860
    draw.line((650, legend_y, 720, legend_y), fill=BLUE, width=7)
    draw.text((735, legend_y), 'Dense', font=pil_font(20), fill=BLACK,
              anchor='lm')
    draw.line((950, legend_y, 1020, legend_y), fill=ORANGE, width=7)
    draw.text((1035, legend_y), 'Hybrid 90/10', font=pil_font(20),
              fill=BLACK, anchor='lm')
    save_canvas(image, FIGURES / 'figure_2_retrieval.png')


def figure_bottleneck():
    data = pd.read_json(
        ROOT / 'runs/answers_eval_qwen2.5-3b__bge-small__s256_o50.jsonl',
        lines=True)
    w, h = 1800, 880
    image = Image.new('RGB', (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    center_text(draw, (w / 2, 55), 'Answer Outcomes Depend on Retrieval',
                pil_font(42, bold=True, serif=True), NAVY)
    center_text(draw, (w / 2, 100),
                'Qwen2.5 3B verdict distribution on all 150 questions',
                pil_font(22), SLATE)
    colors = {'correct': GREEN, 'partial': ORANGE, 'incorrect': RED}
    y_positions = [310, 555]
    labels = [('Gold page missed', False), ('Gold page retrieved', True)]
    bar_x1, bar_x2, bar_h = 435, 1600, 110
    for (label, flag), y in zip(labels, y_positions):
        subset = data[data.page_hit == flag]
        counts = subset.verdict.value_counts()
        total = len(subset)
        draw.text((110, y + bar_h / 2), f'{label}\n(n={total})',
                  font=pil_font(25, bold=True), fill=NAVY, anchor='lm',
                  spacing=6)
        x = bar_x1
        for verdict in ['correct', 'partial', 'incorrect']:
            n = int(counts.get(verdict, 0))
            width = (n / total) * (bar_x2 - bar_x1)
            draw.rectangle((x, y, x + width, y + bar_h),
                           fill=colors[verdict])
            if width > 90:
                center_text(draw, (x + width / 2, y + bar_h / 2),
                            f'{n} ({n/total:.1%})', pil_font(21, bold=True),
                            WHITE)
            x += width
    legend_x = 570
    for verdict in ['correct', 'partial', 'incorrect']:
        draw.rectangle((legend_x, 760, legend_x + 32, 792),
                       fill=colors[verdict])
        draw.text((legend_x + 44, 776), verdict.title(),
                  font=pil_font(20), fill=BLACK, anchor='lm')
        legend_x += 260
    save_canvas(image, FIGURES / 'figure_3_bottleneck.png')


def figure_chunking():
    summary = pd.read_csv(ROOT / 'runs/chunking_pilot_summary.csv')
    labels = ['baseline', 'size384', 'overlap0', 'fixed_windows']
    display = ['Recursive\n256/50', 'Recursive\n384/50',
               'Recursive\n256/0', 'Fixed\n256/50']
    colors = [BLUE, TEAL, ORANGE, SLATE]
    counts = [13797, 8886, 11775, 12849]
    metric = {}
    for label in labels:
        block = summary[summary.variant == label].set_index('metric')
        metric[label] = {
            'MRR': float(block.loc['mrr', 'value']),
            'Evidence hit@10': float(block.loc['evidence_hit@10', 'value']),
        }
    w, h = 1800, 980
    image = Image.new('RGB', (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    center_text(draw, (w / 2, 55), 'Chunking Pilot: Quality and Index Size',
                pil_font(42, bold=True, serif=True), NAVY)
    center_text(draw, (w / 2, 100),
                'Paired 30-question experiment over 26 documents',
                pil_font(22), SLATE)

    left = (90, 220, 1080, 820)
    x1, y1, x2, y2 = left
    draw.text((x1, y1 - 55), 'Retrieval metrics', font=pil_font(28, bold=True),
              fill=NAVY)
    plot_y2 = y2 - 90
    draw.line((x1 + 80, y1, x1 + 80, plot_y2), fill=SLATE, width=3)
    draw.line((x1 + 80, plot_y2, x2, plot_y2), fill=SLATE, width=3)
    for i in range(7):
        val = i / 10
        y = plot_y2 - (val / 0.6) * (plot_y2 - y1)
        draw.line((x1 + 80, y, x2, y), fill='#D9E1E7', width=2)
        draw.text((x1 + 62, y), f'{val:.1f}', font=pil_font(16),
                  fill=SLATE, anchor='rm')
    group_width = (x2 - (x1 + 100)) / len(labels)
    for i, (label, name) in enumerate(zip(labels, display)):
        gx = x1 + 120 + i * group_width
        vals = [metric[label]['MRR'], metric[label]['Evidence hit@10']]
        for j, (v, c) in enumerate(zip(vals, [BLUE, ORANGE])):
            bx1 = gx + j * 65
            by = plot_y2 - (v / 0.6) * (plot_y2 - y1)
            draw.rectangle((bx1, by, bx1 + 50, plot_y2), fill=c)
            draw.text((bx1 + 25, by - 12), f'{v:.2f}',
                      font=pil_font(15, bold=True), fill=c, anchor='ms')
        draw.multiline_text((gx + 58, plot_y2 + 16), name,
                            font=pil_font(16), fill=BLACK, anchor='ma',
                            align='center', spacing=2)
    draw.rectangle((260, 880, 290, 910), fill=BLUE)
    draw.text((305, 895), 'MRR', font=pil_font(18), fill=BLACK, anchor='lm')
    draw.rectangle((470, 880, 500, 910), fill=ORANGE)
    draw.text((515, 895), 'Evidence hit@10', font=pil_font(18),
              fill=BLACK, anchor='lm')

    rx1, ry1, rx2, ry2 = 1190, 220, 1710, 820
    draw.text((rx1, ry1 - 55), 'Index size', font=pil_font(28, bold=True),
              fill=NAVY)
    max_count = 15000
    bar_w = 75
    usable = ry2 - ry1 - 100
    for i, (name, count, color) in enumerate(zip(display, counts, colors)):
        bx = rx1 + 25 + i * 118
        by = ry2 - 90 - (count / max_count) * usable
        draw.rectangle((bx, by, bx + bar_w, ry2 - 90), fill=color)
        draw.text((bx + bar_w / 2, by - 12), f'{count/1000:.1f}k',
                  font=pil_font(16, bold=True), fill=color, anchor='ms')
        draw.multiline_text((bx + bar_w / 2, ry2 - 72), name,
                            font=pil_font(14), fill=BLACK, anchor='ma',
                            align='center', spacing=1)
    save_canvas(image, FIGURES / 'figure_4_chunking.png')


def figure_generator():
    operational = pd.read_csv(
        ROOT / 'runs/generator_size_operational_summary.csv')
    metrics = [
        ('refusal_rate', 'Refusal'),
        ('citation_rate', 'Any citation'),
        ('gold_page_citation_rate', 'Gold-page citation'),
    ]
    w, h = 1800, 900
    image = Image.new('RGB', (w, h), WHITE)
    draw = ImageDraw.Draw(image)
    center_text(draw, (w / 2, 55), 'Generator Size: Grounding Behavior',
                pil_font(42, bold=True, serif=True), NAVY)
    center_text(draw, (w / 2, 100),
                'Matched evaluation on 46 questions with a retrieved gold page',
                pil_font(22), SLATE)
    px1, py1, px2, py2 = 170, 210, 1650, 745
    draw.line((px1, py1, px1, py2), fill=SLATE, width=3)
    draw.line((px1, py2, px2, py2), fill=SLATE, width=3)
    for i in range(6):
        v = i * 0.2
        y = py2 - v * (py2 - py1)
        draw.line((px1, y, px2, y), fill='#D9E1E7', width=2)
        draw.text((px1 - 18, y), f'{v:.0%}', font=pil_font(17),
                  fill=SLATE, anchor='rm')
    group_width = (px2 - px1) / len(metrics)
    for i, (column, label) in enumerate(metrics):
        gx = px1 + i * group_width + 110
        for j, (model, color) in enumerate(zip(['qwen2.5:3b', 'qwen2.5:7b'],
                                               [BLUE, ORANGE])):
            value = float(operational.loc[operational.model == model,
                                          column].iloc[0])
            bx = gx + j * 130
            by = py2 - value * (py2 - py1)
            draw.rectangle((bx, by, bx + 90, py2), fill=color)
            draw.text((bx + 45, by - 12), f'{value:.1%}',
                      font=pil_font(18, bold=True), fill=color, anchor='ms')
        center_text(draw, (gx + 110, py2 + 42), label,
                    pil_font(20, bold=True), BLACK, anchor='ma')
    draw.rectangle((610, 820, 645, 855), fill=BLUE)
    draw.text((660, 838), 'Qwen2.5 3B', font=pil_font(20), fill=BLACK,
              anchor='lm')
    draw.rectangle((950, 820, 985, 855), fill=ORANGE)
    draw.text((1000, 838), 'Qwen2.5 7B', font=pil_font(20), fill=BLACK,
              anchor='lm')
    save_canvas(image, FIGURES / 'figure_5_generator.png')


def generate_figures():
    FIGURES.mkdir(parents=True, exist_ok=True)
    figure_pipeline()
    figure_retrieval()
    figure_bottleneck()
    figure_chunking()
    figure_generator()
    print(f'wrote 5 figures -> {FIGURES}')


def set_run_font(run, name='Arial', size=None, color=BLACK,
                 bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        tc_pr.append(shading)
    shading.set(qn('w:fill'), color.lstrip('#'))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement('w:tblHeader')
    marker.set(qn('w:val'), 'true')
    tr_pr.append(marker)


def set_keep_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement('w:cantSplit')
    tr_pr.append(marker)


def paragraph_bottom_border(paragraph, color='B8C8D3', size='8'):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn('w:pBdr'))
    if p_bdr is None:
        p_bdr = OxmlElement('w:pBdr')
        p_pr.append(p_bdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t')
    text.text = '1'
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=SLATE)


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = False


def add_running_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False

    for header in (section.header, section.even_page_header):
        hp = header.paragraphs[0]
        hp.clear()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(3)
        run = hp.add_run('FINANCEBENCH RAG  |  SCIENTIFIC REPORT')
        set_run_font(run, size=8.5, color=SLATE, bold=True)
        paragraph_bottom_border(hp)

    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(2)
        prefix = fp.add_run('Nayeli  |  EPITA NLP Project A  |  Page ')
        set_run_font(prefix, size=9, color=SLATE)
        add_page_field(fp)


def apply_styles(doc: Document):
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True
    configure_page(section)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.18
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ('Heading 1', 16, NAVY, 16, 8),
        ('Heading 2', 13, BLUE, 12, 6),
        ('Heading 3', 11.5, TEAL, 9, 4),
    ]:
        style = styles[style_name]
        style.font.name = 'Arial'
        style._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
        style._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
    styles['Heading 1'].paragraph_format.page_break_before = False

    caption = styles['Caption']
    caption.font.name = 'Arial'
    caption._element.rPr.rFonts.set(qn('w:ascii'), 'Arial')
    caption._element.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = rgb(SLATE)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.core_properties.title = (
        'Building and Evaluating a Local RAG System for Financial QA')
    doc.core_properties.subject = 'NLP Graded Project A - FinanceBench'
    doc.core_properties.author = 'Nayeli'
    doc.core_properties.keywords = (
        'RAG, FinanceBench, information retrieval, financial question answering')


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('NLP GRADED PROJECT A')
    set_run_font(r, name='Arial', size=11, color=ORANGE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'Building and Evaluating a Local\n'
        'Retrieval-Augmented Generation System')
    set_run_font(r, name='Georgia', size=27, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(42)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Financial Question Answering with FinanceBench')
    set_run_font(r, size=14, color=TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Nayeli')
    set_run_font(r, size=12, color=BLACK, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run('EPITA - Natural Language Processing')
    set_run_font(r, size=10.5, color=SLATE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('July 2026')
    set_run_font(r, size=10.5, color=SLATE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(74)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        'A reproducible, closed-corpus RAG pipeline using free local models')
    set_run_font(r, size=9.5, color=SLATE, italic=True)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section)
    add_running_header_footer(body_section)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`.+?`|\*.+?\*)')


def add_inline(paragraph, text: str):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=10.5)
        token = match.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=10.5, bold=True)
        elif token.startswith('`'):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name='Courier New', size=9, color=BLUE)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=10.5, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=10.5)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    if text.startswith('[') and re.match(r'\[\d+\]', text):
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(5)
        for run in p.runs:
            set_run_font(run, size=9)
    add_inline(p, text)
    if text.startswith('[') and re.match(r'\[\d+\]', text):
        for run in p.runs:
            set_run_font(run, size=9)
    return p


def table_weights(headers: list[str], rows: list[list[str]]) -> list[float]:
    n = len(headers)
    if n == 2:
        if headers[0].lower() == 'stage':
            return [1.6, 4.9]
        if headers[0].lower() == 'component':
            return [2.0, 4.5]
        return [4.5, 2.0]
    if n == 3:
        return [2.1, 1.8, 2.6]
    if n == 4:
        return [2.7, 1.1, 1.4, 1.3]
    if n == 5:
        return [2.4, 1.0, 1.0, 1.0, 1.1]
    if n == 6:
        return [2.0, 0.85, 0.9, 0.9, 0.95, 0.9]
    if n == 7:
        return [1.7, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8]
    return [1.0] * n


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    header = table.rows[0]
    set_repeat_table_header(header)
    set_keep_together(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_fill(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                       else WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(text)
        set_run_font(r, size=8.3, color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        set_keep_together(row)
        for i, text in enumerate(values):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx % 2:
                set_cell_fill(cell, PALE)
            p = cell.paragraphs[0]
            is_number = bool(re.fullmatch(r'[-+]?\d[\d,]*(?:\.\d+)?%?', text))
            p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if is_number or i > 0
                           else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.05
            default_bold = i == 0 and len(text) < 40
            pos = 0
            for match in INLINE_RE.finditer(text):
                if match.start() > pos:
                    r = p.add_run(text[pos:match.start()])
                    set_run_font(r, size=8.0, color=BLACK,
                                 bold=default_bold)
                token = match.group(0)
                if token.startswith('`'):
                    r = p.add_run(token[1:-1])
                    set_run_font(r, name='Courier New', size=7.3,
                                 color=BLUE, bold=False)
                elif token.startswith('**'):
                    r = p.add_run(token[2:-2])
                    set_run_font(r, size=8.0, color=BLACK, bold=True)
                else:
                    r = p.add_run(token[1:-1])
                    set_run_font(r, size=8.0, color=BLACK,
                                 bold=default_bold, italic=True)
                pos = match.end()
            if pos < len(text):
                r = p.add_run(text[pos:])
                set_run_font(r, size=8.0, color=BLACK,
                             bold=default_bold)
    widths = column_widths_from_weights(
        table_weights(headers, rows), total_width_dxa=9360)
    apply_table_geometry(
        table, widths, table_width_dxa=9360, indent_dxa=120,
        cell_margins_dxa={'top': 90, 'bottom': 90,
                          'start': 120, 'end': 120})
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    return table


def add_figure(doc: Document, filename: str, caption: str):
    path = FIGURES / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.2))
    doc_pr = run._r.xpath('.//wp:docPr')[0]
    doc_pr.set('title', filename.replace('_', ' ').removesuffix('.png'))
    doc_pr.set('descr', caption)
    cp = doc.add_paragraph(style='Caption')
    cp.add_run(caption)


def parse_table(lines: list[str], start: int):
    headers = [c.strip() for c in lines[start].strip().strip('|').split('|')]
    rows = []
    i = start + 2
    while i < len(lines) and lines[i].strip().startswith('|'):
        rows.append(
            [c.strip() for c in lines[i].strip().strip('|').split('|')])
        i += 1
    return headers, rows, i


def build_docx():
    doc = Document()
    apply_styles(doc)
    add_cover(doc)
    lines = SOURCE.read_text().splitlines()
    i = 0
    paragraph_parts: list[str] = []

    def flush():
        nonlocal paragraph_parts
        if paragraph_parts:
            add_body_paragraph(doc, ' '.join(part.strip()
                                              for part in paragraph_parts))
            paragraph_parts = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            flush()
            i += 1
            continue
        if stripped.startswith('[[FIGURE:') and stripped.endswith(']]'):
            flush()
            content = stripped[len('[[FIGURE:'):-2]
            filename, caption = content.split('|', 1)
            add_figure(doc, filename, caption)
            i += 1
            continue
        if stripped.startswith('|') and i + 1 < len(lines):
            flush()
            headers, rows, i = parse_table(lines, i)
            add_table(doc, headers, rows)
            continue
        if stripped.startswith('### '):
            flush()
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            flush()
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            flush()
            title = stripped[2:]
            p = doc.add_heading(title, level=1)
            p.paragraph_format.page_break_before = title in {
                '1. Introduction',
                'Appendix A. Reproducibility and artifacts',
            }
            i += 1
            continue
        paragraph_parts.append(stripped)
        i += 1
    flush()

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f'wrote report -> {OUTPUT}')


def main():
    generate_figures()
    build_docx()


if __name__ == '__main__':
    main()
